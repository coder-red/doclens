from __future__ import annotations

import io
import json
import random
from pathlib import Path

from PIL import Image, ImageFilter
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as rl_canvas

FIXTURE_DIR = Path(__file__).resolve().parent


GROUND_TRUTH_A = {
    "file": "layout_a_classic_invoice.pdf",
    "document_type": "invoice",
    "vendor_name": "Meridian Office Supplies Ltd.",
    "invoice_number": "INV-2026-0847",
    "document_date": "2026-07-14",
    "due_date": "2026-08-13",
    "currency": "USD",
    "subtotal": 1980.50,
    "tax_amount": 168.34,
    "total_amount": 2148.84,
}

GROUND_TRUTH_B = {
    "file": "layout_b_receipt_style.pdf",
    "document_type": "receipt",
    "vendor_name": "Café Aroma",
    "invoice_number": "R-1188",
    "document_date": "2026-07-21",
    "currency": "EUR",
    "subtotal": 21.70,
    "tax_amount": 1.50,
    "total_amount": 23.20,
}

GROUND_TRUTH_C = {
    "file": "layout_c_word_invoice.docx",
    "document_type": "invoice",
    "vendor_name": "Northgate Plumbing Services LLC",
    "invoice_number": "NP-26-0311",
    "document_date": "2026-08-02",
    "due_date": "2026-09-01",
    "currency": "USD",
    "subtotal": 845.00,
    "tax_amount": 59.15,
    "total_amount": 904.15,
}


def _draw_layout_a(path: Path) -> None:
    c = rl_canvas.Canvas(str(path), pagesize=letter)
    w, h = letter
    c.setFont("Helvetica-Bold", 20)
    c.drawString(20 * mm, h - 25 * mm, "MERIDIAN OFFICE SUPPLIES LTD.")
    c.setFont("Helvetica", 9)
    c.drawString(20 * mm, h - 31 * mm, "412 Commerce Blvd, Suite 300, Chicago, IL 60601")
    c.drawString(20 * mm, h - 35.5 * mm, "VAT ID: US-47-2913856")

    c.setFont("Helvetica-Bold", 12)
    c.drawRightString(w - 20 * mm, h - 25 * mm, "INVOICE")
    c.setFont("Helvetica", 10)
    c.drawRightString(w - 20 * mm, h - 32 * mm, "Invoice #: INV-2026-0847")
    c.drawRightString(w - 20 * mm, h - 37 * mm, "Date: 14 July 2026")
    c.drawRightString(w - 20 * mm, h - 42 * mm, "Due: 13 August 2026")

    c.setFont("Helvetica-Bold", 10)
    c.drawString(20 * mm, h - 58 * mm, "Bill to:")
    c.setFont("Helvetica", 10)
    c.drawString(20 * mm, h - 64 * mm, "Brightline Consulting LLC")
    c.drawString(20 * mm, h - 69 * mm, "88 Harbor St, Boston, MA 02110")

    rows = [
        ("Executive Mesh Chair", "4", "189.00", "756.00"),
        ("Electric Standing Desk 160cm", "2", "449.50", "899.00"),
        ("Dual Monitor Arm", "6", "54.25", "325.50"),
    ]
    top = h - 90 * mm
    c.setFont("Helvetica-Bold", 9)
    c.drawString(20 * mm, top, "Description")
    c.drawRightString(w - 95 * mm, top, "Qty")
    c.drawRightString(w - 60 * mm, top, "Unit (USD)")
    c.drawRightString(w - 20 * mm, top, "Amount (USD)")
    c.line(20 * mm, top - 2 * mm, w - 20 * mm, top - 2 * mm)

    y = top - 9 * mm
    c.setFont("Helvetica", 9)
    for desc, qty, unit, amount in rows:
        c.drawString(20 * mm, y, desc)
        c.drawRightString(w - 95 * mm, y, qty)
        c.drawRightString(w - 60 * mm, y, unit)
        c.drawRightString(w - 20 * mm, y, amount)
        y -= 7 * mm

    y -= 6 * mm
    c.setFont("Helvetica", 10)
    c.drawRightString(w - 60 * mm, y, "Subtotal:")
    c.drawRightString(w - 20 * mm, y, "1980.50")
    y -= 6 * mm
    c.drawRightString(w - 60 * mm, y, "Sales tax (8.5%):")
    c.drawRightString(w - 20 * mm, y, "168.34")
    y -= 7 * mm
    c.setFont("Helvetica-Bold", 11)
    c.drawRightString(w - 60 * mm, y, "TOTAL DUE:")
    c.drawRightString(w - 20 * mm, y, "$2,148.84")

    c.setFont("Helvetica-Oblique", 8)
    c.drawString(20 * mm, 22 * mm, "Payment via ACH to Meridian Office Supplies Ltd. Reference the invoice number.")
    c.showPage()
    c.save()


def _draw_layout_b(path: Path) -> None:
    c = rl_canvas.Canvas(str(path), pagesize=(80 * mm, 200 * mm))
    w = 80 * mm
    cx = w / 2

    c.setFont("Helvetica-Bold", 15)
    c.drawCentredString(cx, 185 * mm, "Café Aroma")
    c.setFont("Helvetica", 8)
    c.drawCentredString(cx, 179 * mm, "Marktplatz 9, 04109 Leipzig")
    c.drawCentredString(cx, 174.5 * mm, "USt-IdNr: DE 812 345 678")

    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(cx, 165 * mm, "RECEIPT R-1188")
    c.setFont("Helvetica", 8)
    c.drawCentredString(cx, 160 * mm, "21.07.2026 14:32")

    items = [
        ("Flat White x2 @ 3.80", "7.60"),
        ("Almond Croissant", "4.20"),
        ("Club Sandwich", "9.90"),
    ]
    y = 150 * mm
    for desc, amount in items:
        c.setFont("Helvetica", 9)
        c.drawString(8 * mm, y, desc)
        c.drawRightString(w - 8 * mm, y, amount)
        y -= 6.5 * mm

    y -= 3 * mm
    c.line(8 * mm, y, w - 8 * mm, y)
    y -= 6.5 * mm
    c.setFont("Helvetica", 9)
    c.drawString(8 * mm, y, "Subtotal")
    c.drawRightString(w - 8 * mm, y, "EUR 21.70")
    y -= 6 * mm
    c.drawString(8 * mm, y, "VAT 7%")
    c.drawRightString(w - 8 * mm, y, "EUR 1.50")
    y -= 7.5 * mm
    c.setFont("Helvetica-Bold", 11)
    c.drawString(8 * mm, y, "TOTAL")
    c.drawRightString(w - 8 * mm, y, "23.20")

    y -= 10 * mm
    c.setFont("Helvetica-Oblique", 8)
    c.drawCentredString(cx, y, "* card payment * contactless *")
    c.drawCentredString(cx, y - 5 * mm, "Danke! Tschüss!")
    c.showPage()
    c.save()


def _build_layout_c_docx(path: Path) -> None:
    import docx
    from docx.shared import Pt

    d = docx.Document()
    d.add_heading("Northgate Plumbing Services LLC", level=1)
    p = d.add_paragraph("87 Ridgeway Ave, Portland, OR 97205\nLicense #OR-PLB-44120")
    p.runs[0].font.size = Pt(9)
    d.add_paragraph("")
    meta = d.add_paragraph()
    meta.add_run("INVOICE NP-26-0311\n").bold = True
    meta.add_run("Invoice date: August 2, 2026\nDue date: September 1, 2026\nBill to: Hale Residence, 12 Cedar Ct")
    d.add_paragraph("")

    table = d.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    rows = [
        ("Description", "Qty", "Rate (USD)", "Amount (USD)"),
        ("Emergency leak repair", "2.5", "120.00", "300.00"),
        ("Replace angle stop valves", "4", "38.75", "155.00"),
        ("Water heater flush", "1", "190.00", "190.00"),
        ("Pipe insulation (per ft)", "50", "4.00", "200.00"),
    ]
    for r_idx, row_vals in enumerate(rows):
        for c_idx, val in enumerate(row_vals):
            para = table.rows[r_idx].cells[c_idx].paragraphs[0]
            run = para.add_run(val)
            run.bold = r_idx == 0

    d.add_paragraph("")
    totals = d.add_paragraph()
    totals.add_run("Subtotal: $845.00\n").bold = False
    totals.add_run("Tax (7%): $59.15\n")
    totals.add_run("TOTAL DUE: $904.15").bold = True
    d.add_paragraph("")
    terms = d.add_paragraph("Payment due within 30 days. Make checks to Northgate Plumbing Services LLC.")
    terms.runs[0].font.size = Pt(8)
    d.save(str(path))


def _degraded_variant(source_pdf: Path, target_png: Path) -> None:
    import pymupdf

    doc = pymupdf.open(source_pdf)
    pix = doc[0].get_pixmap(matrix=pymupdf.Matrix(1.6, 1.6))
    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    doc.close()
    img = img.rotate(2.4, resample=Image.BICUBIC, expand=True, fillcolor=(210, 205, 195))
    noise = Image.effect_noise(img.size, 28).convert("RGB")
    img = Image.blend(img, noise, 0.22)
    img = img.filter(ImageFilter.GaussianBlur(0.8))
    img.save(target_png, format="JPEG", quality=38)


def generate_all(force: bool = False) -> dict[str, Path]:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    outputs: dict[str, Path] = {}

    pdf_a = FIXTURE_DIR / GROUND_TRUTH_A["file"]
    if force or not pdf_a.exists():
        _draw_layout_a(pdf_a)
    outputs["classic_invoice"] = pdf_a

    pdf_b = FIXTURE_DIR / GROUND_TRUTH_B["file"]
    if force or not pdf_b.exists():
        _draw_layout_b(pdf_b)
    outputs["receipt_style"] = pdf_b

    docx_c = FIXTURE_DIR / GROUND_TRUTH_C["file"]
    if force or not docx_c.exists():
        _build_layout_c_docx(docx_c)
    outputs["word_invoice"] = docx_c

    degraded = FIXTURE_DIR / "layout_b_receipt_degraded.jpg"
    if force or not degraded.exists():
        _degraded_variant(pdf_b, degraded)
    outputs["degraded_receipt"] = degraded

    gt_path = FIXTURE_DIR / "ground_truth.json"
    gt = {"layout_a": GROUND_TRUTH_A, "layout_b": GROUND_TRUTH_B, "layout_c": GROUND_TRUTH_C}
    gt_path.write_text(json.dumps(gt, indent=2, ensure_ascii=False), encoding="utf-8")
    outputs["ground_truth"] = gt_path
    return outputs


if __name__ == "__main__":
    paths = generate_all(force=True)
    for name, p in paths.items():
        print(f"{name}: {p}")
